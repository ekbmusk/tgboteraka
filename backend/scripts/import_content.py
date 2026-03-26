"""
Import physics content from Word (.docx) files into the database.

Usage (from backend/ directory):
    python -m scripts.import_content tests
    python -m scripts.import_content problems
    python -m scripts.import_content lectures
    python -m scripts.import_content all
"""

import argparse
import os
import re
import sys
from pathlib import Path

from docx import Document

# Ensure app modules are importable
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from dotenv import load_dotenv

load_dotenv(Path(__file__).resolve().parent.parent.parent / ".env")

from app.database.database import SessionLocal, create_tables
from app.models.admin_test import AdminTestQuestion
from app.models.problem import Problem
from app.models.theory_content import TheoryContent

DATA_DIR = Path(__file__).resolve().parent.parent.parent / "Data"

# Topic names for each of the 15 test files (mapped by file number)
TEST_TOPIC_MAP = {
    1: "Кинематика негіздері",
    2: "Жылдамдық және үдеу",
    3: "Ньютон заңдары (1)",
    4: "Ньютон заңдары (2)",
    5: "Жұмыс, энергия, қуат",
    6: "Сақталу заңдары",
    7: "Айналмалы қозғалыс",
    8: "Импульс моменті",
    9: "Бүкіл әлемдік тартылыс",
    10: "Салыстырмалылық теориясы",
    11: "Қысым және гидростатика",
    12: "Бернулли теңдеуі",
    13: "Сұйықтық динамикасы",
    14: "Тербелістер",
    15: "Акустика",
}

# Unicode subscript/superscript maps for LaTeX conversion
SUPERSCRIPTS = {"⁰": "0", "¹": "1", "²": "2", "³": "3", "⁴": "4",
                "⁵": "5", "⁶": "6", "⁷": "7", "⁸": "8", "⁹": "9",
                "⁺": "+", "⁻": "-", "ⁿ": "n"}
SUBSCRIPTS = {"₀": "0", "₁": "1", "₂": "2", "₃": "3", "₄": "4",
              "₅": "5", "₆": "6", "₇": "7", "₈": "8", "₉": "9",
              "ₓ": "x", "ₙ": "n"}
GREEK_MAP = {
    "α": "\\alpha", "β": "\\beta", "γ": "\\gamma", "δ": "\\delta",
    "ε": "\\epsilon", "θ": "\\theta", "λ": "\\lambda", "μ": "\\mu",
    "ν": "\\nu", "π": "\\pi", "ρ": "\\rho", "σ": "\\sigma",
    "τ": "\\tau", "φ": "\\varphi", "ω": "\\omega",
    "Δ": "\\Delta", "Σ": "\\Sigma", "Ω": "\\Omega",
}


def plain_to_latex(text: str) -> str:
    """Best-effort conversion of plain-text formula to LaTeX."""
    result = text.strip()

    # Handle superscripts: v² → v^{2}
    for char, val in SUPERSCRIPTS.items():
        result = result.replace(char, f"^{{{val}}}")

    # Handle subscripts: x₁ → x_{1}
    for char, val in SUBSCRIPTS.items():
        result = result.replace(char, f"_{{{val}}}")

    # Handle Greek letters
    for char, latex in GREEK_MAP.items():
        result = result.replace(char, latex)

    # Handle vector arrows: r⃗ → \vec{r}
    result = re.sub(r"(\w)\u20d7", r"\\vec{\1}", result)

    return f"${result}$"


def is_formula_paragraph(text: str) -> bool:
    """Heuristic: detect if a paragraph is a formula rather than prose."""
    text = text.strip()
    if not text or len(text) > 120:
        return False
    if "=" not in text:
        return False
    # Short text with = sign and few words → likely formula
    words = text.split()
    if len(words) <= 8:
        return True
    return False


def is_problem_marker(text: str) -> bool:
    """Check if paragraph is a problem number marker like '1️⃣ Есеп'."""
    stripped = text.strip()
    if "Есеп" in stripped and len(stripped) < 30:
        return True
    return False


def extract_topic_from_header(text: str) -> str:
    """Extract topic name from section header, looking for «...» quotes."""
    match = re.search(r"[«\"](.*?)[»\"]", text)
    if match:
        topic = match.group(1)
        # Truncate if too long
        if len(topic) > 120:
            topic = topic[:117] + "..."
        return topic
    # Fallback: use first 80 chars
    return text[:80].strip()


# ─── TEST IMPORT ──────────────────────────────────────────────────────────────


def parse_test_file(filepath: str) -> list[dict]:
    """Parse a single test .docx file and return list of question dicts."""
    doc = Document(filepath)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    questions = []
    answers = {}
    current_level = None
    in_answers = False

    for para_text in paragraphs:
        # Check for answer section
        if "ЖАУАПТАРЫ" in para_text.upper():
            in_answers = True
            continue

        if in_answers:
            # Parse answer lines: "1-B\n2-C\n..."
            # Also skip level labels like "Жеңіл деңгей:"
            for line in para_text.split("\n"):
                line = line.strip()
                match = re.match(r"(\d+)\s*[-–]\s*([A-Da-d])", line)
                if match:
                    num = int(match.group(1))
                    letter = match.group(2).upper()
                    answers[num] = letter
            continue

        # Detect level headers
        upper = para_text.upper()
        if "ЖЕҢІЛ ДЕҢГЕЙ" in upper or "ЖЕҢІЛ" in upper and "ДЕҢГЕЙ" in upper:
            current_level = "easy"
            continue
        if "ОРТА ДЕҢГЕЙ" in upper:
            current_level = "medium"
            continue
        if "КҮРДЕЛІ ДЕҢГЕЙ" in upper:
            current_level = "hard"
            continue

        if not current_level:
            continue

        # Parse question: "Question text\nA) opt\nB) opt\nC) opt\nD) opt"
        lines = para_text.split("\n")
        question_text = None
        options = {}

        for line in lines:
            line = line.strip()
            opt_match = re.match(r"^([A-D])\)\s*(.+)", line)
            if opt_match:
                options[opt_match.group(1)] = opt_match.group(2).strip()
            elif not question_text and line:
                question_text = line

        if question_text and len(options) == 4:
            q_num = len(questions) + 1
            questions.append({
                "question": question_text,
                "options": options,
                "level": current_level,
                "number": q_num,
            })

    # Merge answers
    for q in questions:
        q["correct"] = answers.get(q["number"], "A")

    return questions


def import_tests(db):
    """Import test questions from all 15 test files."""
    test_dir = DATA_DIR / "Тест (механика) үш деңгей" / "1. Тест-жеңіл"
    if not test_dir.exists():
        print(f"Test directory not found: {test_dir}")
        return

    created = 0
    skipped = 0

    for file_num in range(1, 16):
        filename = f"{file_num}. тест оңай.docx"
        filepath = test_dir / filename
        if not filepath.exists():
            print(f"  Skipping missing file: {filename}")
            continue

        topic = TEST_TOPIC_MAP.get(file_num, f"Тақырып {file_num}")
        questions = parse_test_file(str(filepath))

        for q in questions:
            # Check for duplicate
            existing = db.query(AdminTestQuestion).filter(
                AdminTestQuestion.question == q["question"]
            ).first()
            if existing:
                skipped += 1
                continue

            db.add(AdminTestQuestion(
                topic=topic,
                question=q["question"],
                option_a=q["options"].get("A", ""),
                option_b=q["options"].get("B", ""),
                option_c=q["options"].get("C", ""),
                option_d=q["options"].get("D", ""),
                correct_option=q["correct"],
                explanation=None,
            ))
            created += 1

        print(f"  File {file_num}: {len(questions)} questions parsed")

    db.commit()
    print(f"Tests: {created} created, {skipped} skipped (duplicates)")


# ─── PROBLEMS IMPORT ─────────────────────────────────────────────────────────


def import_problems(db):
    """Import problems from the mechanics problems file."""
    problems_file = DATA_DIR / "Есеп (механика) үш деңгей" / "есеп.docx"
    if not problems_file.exists():
        print(f"Problems file not found: {problems_file}")
        return

    doc = Document(str(problems_file))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # Step 1: Find section boundaries
    sections = []
    for i, text in enumerate(paragraphs):
        if text.startswith("Төменде") or text.startswith("🌐"):
            topic = extract_topic_from_header(text)
            sections.append({"start": i, "topic": topic})

    created = 0
    skipped = 0

    # Step 2: Parse each section
    for sec_idx, section in enumerate(sections):
        end = sections[sec_idx + 1]["start"] if sec_idx + 1 < len(sections) else len(paragraphs)
        topic = section["topic"]
        current_level = None
        current_problem_lines = []
        current_formula = None

        def save_problem():
            nonlocal created, skipped, current_problem_lines, current_formula
            if not current_problem_lines or not current_level:
                current_problem_lines = []
                current_formula = None
                return

            question = "\n".join(current_problem_lines)
            # Check duplicate
            existing = db.query(Problem).filter(
                Problem.topic == topic,
                Problem.question == question,
            ).first()
            if existing:
                skipped += 1
            else:
                db.add(Problem(
                    topic=topic,
                    question=question,
                    formula=current_formula,
                    correct_answer="",
                    solution=None,
                    difficulty=current_level,
                    tags=[],
                ))
                created += 1

            current_problem_lines = []
            current_formula = None

        for i in range(section["start"] + 1, end):
            text = paragraphs[i]

            # Detect difficulty markers
            if "🟢" in text:
                save_problem()
                current_level = "easy"
                continue
            if "🟡" in text:
                save_problem()
                current_level = "medium"
                continue
            if "🔴" in text:
                save_problem()
                current_level = "hard"
                continue

            # Detect problem number markers
            if is_problem_marker(text):
                save_problem()
                continue

            if not current_level:
                continue

            # Detect LaTeX formula blocks wrapped in [...]
            if text.startswith("[") and "\\" in text:
                current_formula = text.strip("[] \n")
                continue

            # Accumulate problem text
            if text:
                current_problem_lines.append(text)

        # Save last problem in section
        save_problem()

    db.commit()
    print(f"Problems: {created} created, {skipped} skipped (duplicates)")


# ─── LECTURES IMPORT ─────────────────────────────────────────────────────────


def import_lectures(db):
    """Import lectures as theory content blocks."""
    lectures_dir = DATA_DIR / "Лекция 1-15 апта"
    if not lectures_dir.exists():
        print(f"Lectures directory not found: {lectures_dir}")
        return

    updated = 0
    created = 0

    for num in range(1, 16):
        filename = f"{num}-лекция.docx"
        filepath = lectures_dir / filename
        if not filepath.exists():
            print(f"  Skipping missing file: {filename}")
            continue

        topic_id = f"lecture_{num:02d}"
        doc = Document(str(filepath))

        blocks = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Classify paragraph
            lower = text.lower()
            if lower.startswith("мысал") or lower.startswith("мысалы"):
                blocks.append({"type": "example", "content": text})
            elif is_formula_paragraph(text):
                blocks.append({"type": "formula", "content": plain_to_latex(text)})
            else:
                blocks.append({"type": "text", "content": text})

        # Extract title from first paragraph or file name
        title = f"{num}-лекция"
        if blocks:
            first_text = blocks[0].get("content", "")
            if len(first_text) < 100:
                title = first_text

        # Upsert
        existing = db.query(TheoryContent).filter(
            TheoryContent.topic_id == topic_id
        ).first()
        if existing:
            existing.blocks = blocks
            if existing.title == f"{num}-лекция" or not existing.blocks:
                existing.title = title
            updated += 1
        else:
            db.add(TheoryContent(topic_id=topic_id, title=title, blocks=blocks))
            created += 1

        print(f"  Lecture {num}: {len(blocks)} blocks")

    db.commit()
    print(f"Lectures: {created} created, {updated} updated")


# ─── MAIN ────────────────────────────────────────────────────────────────────


def main():
    parser = argparse.ArgumentParser(description="Import physics content from Word files")
    parser.add_argument(
        "what",
        choices=["tests", "problems", "lectures", "all"],
        help="What to import",
    )
    args = parser.parse_args()

    # Ensure tables exist
    create_tables()

    db = SessionLocal()
    try:
        if args.what in ("tests", "all"):
            print("=== Importing tests ===")
            import_tests(db)

        if args.what in ("problems", "all"):
            print("=== Importing problems ===")
            import_problems(db)

        if args.what in ("lectures", "all"):
            print("=== Importing lectures ===")
            import_lectures(db)

        print("\nDone!")
    finally:
        db.close()


if __name__ == "__main__":
    main()
